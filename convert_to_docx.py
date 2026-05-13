import docx
import re

md_path = r"C:\Users\Shreyansh\.gemini\antigravity\brain\e0d432b4-87c8-40c3-b3cd-ee92d04e0164\artifacts\AI_Pipeline_Structure.md"
docx_path = r"c:\Users\Shreyansh\OneDrive\Desktop\Advance Editor\AI_Pipeline_Structure.docx"

doc = docx.Document()

# Read markdown
try:
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
except FileNotFoundError:
    print(f"File not found: {md_path}")
    exit(1)

in_code_block = False
code_text = ""

for line in lines:
    line = line.strip('\n')
    if line.startswith('```'):
        if in_code_block:
            doc.add_paragraph(code_text.strip())
            code_text = ""
            in_code_block = False
        else:
            in_code_block = True
        continue
    
    if in_code_block:
        code_text += line + "\n"
        continue
        
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    elif line.startswith('### '):
        doc.add_heading(line[4:], level=3)
    elif line.startswith('- '):
        p = doc.add_paragraph(style='List Bullet')
        text = line[2:]
        # Replace inline code with quotes
        text = text.replace('`', "'")
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)
    elif line.strip() == "":
        pass
    else:
        p = doc.add_paragraph()
        text = line
        text = text.replace('`', "'")
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                p.add_run(part[2:-2]).bold = True
            else:
                p.add_run(part)

doc.save(docx_path)
print(f"Successfully saved to {docx_path}")
